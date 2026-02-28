"""
报告生成器模块
自动生成标准化的审核报告，支持多种格式和自定义模板
"""

import pandas as pd
import numpy as np
import logging
from typing import Dict, List, Any, Optional, Union
from datetime import datetime, timedelta
from pathlib import Path
from dataclasses import dataclass, field
import json
from jinja2 import Environment, FileSystemLoader, Template
import base64
import io

# 尝试导入可选依赖
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab not available. PDF generation will be disabled.")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Word generation will be disabled.")

try:
    import matplotlib.pyplot as plt
    import seaborn as sns
    import plotly.graph_objects as go
    import plotly.express as px
    from plotly.subplots import make_subplots
    import plotly.io as pio
    PLOTTING_AVAILABLE = True
    # 设置中文字体
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS']
    plt.rcParams['axes.unicode_minus'] = False
except ImportError:
    PLOTTING_AVAILABLE = False
    logging.warning("Plotting libraries not available. Charts will be disabled.")

logger = logging.getLogger(__name__)


@dataclass
class ReportConfig:
    """报告配置"""
    title: str = "账务审核报告"
    author: str = "Learn Accounting Agent"
    company: str = "示例公司"
    report_period: str = ""
    include_charts: bool = True
    include_details: bool = True
    include_recommendations: bool = True
    language: str = "zh"  # zh, en
    template_dir: str = "templates"
    output_dir: str = "reports"
    chart_style: str = "seaborn"  # seaborn, plotly, matplotlib
    chart_format: str = "png"  # png, svg, pdf


@dataclass
class ReportData:
    """报告数据"""
    summary: Dict[str, Any] = field(default_factory=dict)
    audit_results: List[Dict[str, Any]] = field(default_factory=list)
    risk_analysis: Dict[str, Any] = field(default_factory=dict)
    trend_analysis: Dict[str, Any] = field(default_factory=dict)
    recommendations: List[str] = field(default_factory=list)
    details: pd.DataFrame = field(default_factory=pd.DataFrame)
    charts: Dict[str, str] = field(default_factory=dict)  # chart_name -> file_path


class ChartGenerator:
    """图表生成器"""
    
    def __init__(self, config: ReportConfig):
        self.config = config
        self.chart_dir = Path(config.output_dir) / "charts"
        self.chart_dir.mkdir(parents=True, exist_ok=True)
        
    def generate_risk_distribution_chart(self, risk_data: Dict[str, Any]) -> str:
        """生成风险分布图"""
        if not PLOTTING_AVAILABLE:
            return ""
            
        try:
            risk_counts = risk_data.get("risk_distribution", {})
            if not risk_counts:
                return ""
                
            if self.config.chart_style == "plotly":
                return self._generate_plotly_pie_chart(risk_counts, "risk_distribution")
            else:
                return self._generate_matplotlib_pie_chart(risk_counts, "risk_distribution")
                
        except Exception as e:
            logger.error(f"生成风险分布图失败: {e}")
            return ""
            
    def generate_trend_chart(self, trend_data: Dict[str, Any]) -> str:
        """生成趋势图"""
        if not PLOTTING_AVAILABLE:
            return ""
            
        try:
            daily_trends = trend_data.get("daily_trends", {})
            if not daily_trends:
                return ""
                
            # 转换数据格式
            dates = list(daily_trends.keys())
            total_counts = [daily_trends[date]["total"] for date in dates]
            failed_counts = [daily_trends[date]["failed"] for date in dates]
            
            if self.config.chart_style == "plotly":
                return self._generate_plotly_trend_chart(dates, total_counts, failed_counts)
            else:
                return self._generate_matplotlib_trend_chart(dates, total_counts, failed_counts)
                
        except Exception as e:
            logger.error(f"生成趋势图失败: {e}")
            return ""
            
    def generate_task_performance_chart(self, task_data: Dict[str, Any]) -> str:
        """生成任务性能图"""
        if not PLOTTING_AVAILABLE:
            return ""
            
        try:
            task_stats = task_data.get("task_distribution", {})
            if not task_stats:
                return ""
                
            tasks = list(task_stats.keys())
            counts = [task_stats[task] for task in tasks]
            
            if self.config.chart_style == "plotly":
                return self._generate_plotly_bar_chart(tasks, counts, "task_performance")
            else:
                return self._generate_matplotlib_bar_chart(tasks, counts, "task_performance")
                
        except Exception as e:
            logger.error(f"生成任务性能图失败: {e}")
            return ""
            
    def _generate_plotly_pie_chart(self, data: Dict[str, str], chart_name: str) -> str:
        """生成Plotly饼图"""
        fig = go.Figure(data=[go.Pie(
            labels=list(data.keys()),
            values=list(data.values()),
            hole=0.3,
            textinfo='label+percent',
            textfont_size=12
        )])
        
        fig.update_layout(
            title="风险等级分布",
            font=dict(size=12),
            showlegend=True
        )
        
        chart_path = self.chart_dir / f"{chart_name}.{self.config.chart_format}"
        if self.config.chart_format == "html":
            fig.write_html(str(chart_path))
        else:
            fig.write_image(str(chart_path))
            
        return str(chart_path)
        
    def _generate_matplotlib_pie_chart(self, data: Dict[str, str], chart_name: str) -> str:
        """生成Matplotlib饼图"""
        plt.figure(figsize=(8, 6))
        colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99']
        
        plt.pie(
            data.values(),
            labels=data.keys(),
            autopct='%1.1f%%',
            colors=colors[:len(data)],
            startangle=90
        )
        
        plt.title('风险等级分布', fontsize=14, fontweight='bold')
        plt.axis('equal')
        
        chart_path = self.chart_dir / f"{chart_name}.{self.config.chart_format}"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return str(chart_path)
        
    def _generate_plotly_trend_chart(self, dates: List[str], totals: List[int], 
                                    faileds: List[int], chart_name: str) -> str:
        """生成Plotly趋势图"""
        fig = go.Figure()
        
        fig.add_trace(go.Scatter(
            x=dates,
            y=totals,
            mode='lines+markers',
            name='总记录数',
            line=dict(color='blue', width=2)
        ))
        
        fig.add_trace(go.Scatter(
            x=dates,
            y=faileds,
            mode='lines+markers',
            name='失败记录数',
            line=dict(color='red', width=2)
        ))
        
        fig.update_layout(
            title='审核趋势分析',
            xaxis_title='日期',
            yaxis_title='记录数量',
            hovermode='x unified',
            font=dict(size=12)
        )
        
        chart_path = self.chart_dir / f"{chart_name}.{self.config.chart_format}"
        if self.config.chart_format == "html":
            fig.write_html(str(chart_path))
        else:
            fig.write_image(str(chart_path))
            
        return str(chart_path)
        
    def _generate_matplotlib_trend_chart(self, dates: List[str], totals: List[int], 
                                       faileds: List[int], chart_name: str) -> str:
        """生成Matplotlib趋势图"""
        plt.figure(figsize=(12, 6))
        
        x = range(len(dates))
        
        plt.plot(x, totals, 'b-o', label='总记录数', linewidth=2, markersize=6)
        plt.plot(x, faileds, 'r-s', label='失败记录数', linewidth=2, markersize=6)
        
        plt.xlabel('日期')
        plt.ylabel('记录数量')
        plt.title('审核趋势分析', fontsize=14, fontweight='bold')
        plt.legend()
        plt.grid(True, alpha=0.3)
        
        # 设置x轴标签
        plt.xticks(x[::max(1, len(dates)//10)], dates[::max(1, len(dates)//10)], rotation=45)
        
        chart_path = self.chart_dir / f"{chart_name}.{self.config.chart_format}"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return str(chart_path)
        
    def _generate_plotly_bar_chart(self, categories: List[str], values: List[int], 
                                  chart_name: str) -> str:
        """生成Plotly柱状图"""
        fig = go.Figure(data=[
            go.Bar(x=categories, y=values, marker_color='lightblue')
        ])
        
        fig.update_layout(
            title='任务执行统计',
            xaxis_title='任务类型',
            yaxis_title='执行次数',
            font=dict(size=12)
        )
        
        chart_path = self.chart_dir / f"{chart_name}.{self.config.chart_format}"
        if self.config.chart_format == "html":
            fig.write_html(str(chart_path))
        else:
            fig.write_image(str(chart_path))
            
        return str(chart_path)
        
    def _generate_matplotlib_bar_chart(self, categories: List[str], values: List[int], 
                                     chart_name: str) -> str:
        """生成Matplotlib柱状图"""
        plt.figure(figsize=(10, 6))
        
        bars = plt.bar(categories, values, color='lightblue', alpha=0.7)
        
        plt.xlabel('任务类型')
        plt.ylabel('执行次数')
        plt.title('任务执行统计', fontsize=14, fontweight='bold')
        plt.grid(True, alpha=0.3)
        
        # 添加数值标签
        for bar, value in zip(bars, values):
            plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                    str(value), ha='center', va='bottom')
        
        plt.xticks(rotation=45)
        
        chart_path = self.chart_dir / f"{chart_name}.{self.config.chart_format}"
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        
        return str(chart_path)


class ReportGenerator:
    """报告生成器主类"""
    
    def __init__(self, config: Optional[ReportConfig] = None):
        """
        初始化报告生成器
        
        Args:
            config: 报告配置
        """
        self.config = config or ReportConfig()
        self.template_env = self._setup_template_environment()
        self.chart_generator = ChartGenerator(self.config)
        
        # 确保输出目录存在
        Path(self.config.output_dir).mkdir(parents=True, exist_ok=True)
        
    def _setup_template_environment(self) -> Environment:
        """设置模板环境"""
        template_dir = Path(self.config.template_dir)
        if template_dir.exists():
            loader = FileSystemLoader(str(template_dir))
            env = Environment(loader=loader, autoescape=True)
        else:
            # 使用内置模板
            env = Environment(autoescape=True)
            self._add_builtin_templates(env)
            
        return env
        
    def _add_builtin_templates(self, env: Environment):
        """添加内置模板"""
        # HTML报告模板
        html_template = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ title }}</title>
    <style>
        body { font-family: 'Microsoft YaHei', Arial, sans-serif; margin: 20px; line-height: 1.6; }
        .header { text-align: center; border-bottom: 2px solid #333; padding-bottom: 20px; margin-bottom: 30px; }
        .section { margin-bottom: 30px; }
        .chart { text-align: center; margin: 20px 0; }
        .table { border-collapse: collapse; width: 100%; margin: 20px 0; }
        .table th, .table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        .table th { background-color: #f2f2f2; font-weight: bold; }
        .risk-high { color: #d32f2f; }
        .risk-medium { color: #f57c00; }
        .risk-low { color: #388e3c; }
        .summary-item { margin: 10px 0; }
        .recommendation { background-color: #e3f2fd; padding: 15px; border-radius: 5px; margin: 10px 0; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ title }}</h1>
        <p>生成时间: {{ generation_time }}</p>
        <p>报告期间: {{ report_period }}</p>
    </div>
    
    <div class="section">
        <h2>审核概要</h2>
        <div class="summary-item">
            <strong>总记录数:</strong> {{ summary.total_records | default(0) }}
        </div>
        <div class="summary-item">
            <strong>通过记录数:</strong> {{ summary.passed_records | default(0) }}
        </div>
        <div class="summary-item">
            <strong>失败记录数:</strong> {{ summary.failed_records | default(0) }}
        </div>
        <div class="summary-item">
            <strong>通过率:</strong> {{ "%.2f"|format(summary.pass_rate*100) }}%
        </div>
    </div>
    
    {% if charts.risk_distribution %}
    <div class="section">
        <h2>风险分布</h2>
        <div class="chart">
            <img src="{{ charts.risk_distribution }}" alt="风险分布图" style="max-width: 100%;">
        </div>
    </div>
    {% endif %}
    
    {% if charts.trend_analysis %}
    <div class="section">
        <h2>趋势分析</h2>
        <div class="chart">
            <img src="{{ charts.trend_analysis }}" alt="趋势分析图" style="max-width: 100%;">
        </div>
    </div>
    {% endif %}
    
    {% if recommendations %}
    <div class="section">
        <h2>建议措施</h2>
        {% for rec in recommendations %}
        <div class="recommendation">
            {{ loop.index }}. {{ rec }}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if include_details and details is defined %}
    <div class="section">
        <h2>详细记录</h2>
        <table class="table">
            <thead>
                <tr>
                    <th>记录ID</th>
                    <th>审核日期</th>
                    <th>任务类型</th>
                    <th>风险等级</th>
                    <th>状态</th>
                </tr>
            </thead>
            <tbody>
                {% for record in details %}
                <tr>
                    <td>{{ record.record_id }}</td>
                    <td>{{ record.audit_date }}</td>
                    <td>{{ record.task_type }}</td>
                    <td class="risk-{{ record.risk_level }}">{{ record.risk_level }}</td>
                    <td>{{ "通过" if record.passed else "失败" }}</td>
                </tr>
                {% endfor %}
            </tbody>
        </table>
    </div>
    {% endif %}
    
    <div class="section">
        <p><em>报告由 {{ author }} 自动生成</em></p>
    </div>
</body>
</html>
        """
        env.globals['html_template'] = html_template
        
    def generate_html_report(self, data: ReportData) -> str:
        """生成HTML报告"""
        try:
            template = self.template_env.get_template('html_template') if 'html_template' in self.template_env.globals else None
            
            if template is None:
                template = Template(self.template_env.globals['html_template'])
                
            # 准备模板数据
            template_data = {
                'title': self.config.title,
                'author': self.config.author,
                'generation_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'report_period': self.config.report_period,
                'summary': data.summary,
                'charts': data.charts,
                'recommendations': data.recommendations,
                'include_details': self.config.include_details,
                'details': data.details.to_dict('records') if not data.details.empty else []
            }
            
            # 渲染模板
            html_content = template.render(**template_data)
            
            # 保存文件
            report_path = Path(self.config.output_dir) / f"audit_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
            logger.info(f"HTML报告已生成: {report_path}")
            return str(report_path)
            
        except Exception as e:
            logger.error(f"生成HTML报告失败: {e}")
            raise
            
    def generate_pdf_report(self, data: ReportData) -> str:
        """生成PDF报告"""
        if not REPORTLAB_AVAILABLE:
            logger.error("ReportLab未安装，无法生成PDF报告")
            return ""
            
        try:
            report_path = Path(self.config.output_dir) / f"audit_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            # 创建PDF文档
            doc = SimpleDocTemplate(str(report_path), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # 标题
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # 居中
            )
            story.append(Paragraph(self.config.title, title_style))
            story.append(Spacer(1, 12))
            
            # 报告信息
            info_style = styles['Normal']
            story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", info_style))
            story.append(Paragraph(f"报告期间: {self.config.report_period}", info_style))
            story.append(Spacer(1, 20))
            
            # 审核概要
            story.append(Paragraph("审核概要", styles['Heading2']))
            summary = data.summary
            story.append(Paragraph(f"总记录数: {summary.get('total_records', 0)}", info_style))
            story.append(Paragraph(f"通过记录数: {summary.get('passed_records', 0)}", info_style))
            story.append(Paragraph(f"失败记录数: {summary.get('failed_records', 0)}", info_style))
            story.append(Paragraph(f"通过率: {summary.get('pass_rate', 0):.2%}", info_style))
            story.append(Spacer(1, 20))
            
            # 建议措施
            if data.recommendations:
                story.append(Paragraph("建议措施", styles['Heading2']))
                for i, rec in enumerate(data.recommendations, 1):
                    story.append(Paragraph(f"{i}. {rec}", info_style))
                story.append(Spacer(1, 20))
                
            # 详细记录表格
            if self.config.include_details and not data.details.empty:
                story.append(Paragraph("详细记录", styles['Heading2']))
                
                # 准备表格数据
                table_data = [['记录ID', '审核日期', '任务类型', '风险等级', '状态']]
                for _, record in data.details.iterrows():
                    table_data.append([
                        str(record.get('record_id', '')),
                        str(record.get('audit_date', '')),
                        str(record.get('task_type', '')),
                        str(record.get('risk_level', '')),
                        '通过' if record.get('passed', False) else '失败'
                    ])
                
                # 创建表格
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(table)
                
            # 构建PDF
            doc.build(story)
            
            logger.info(f"PDF报告已生成: {report_path}")
            return str(report_path)
            
        except Exception as e:
            logger.error(f"生成PDF报告失败: {e}")
            raise
            
    def generate_word_report(self, data: ReportData) -> str:
        """生成Word报告"""
        if not PYTHON_DOCX_AVAILABLE:
            logger.error("python-docx未安装，无法生成Word报告")
            return ""
            
        try:
            report_path = Path(self.config.output_dir) / f"audit_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            # 创建Word文档
            doc = Document()
            
            # 标题
            title = doc.add_heading(self.config.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 报告信息
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"报告期间: {self.config.report_period}")
            doc.add_paragraph("")  # 空行
            
            # 审核概要
            doc.add_heading("审核概要", level=1)
            summary = data.summary
            doc.add_paragraph(f"总记录数: {summary.get('total_records', 0)}")
            doc.add_paragraph(f"通过记录数: {summary.get('passed_records', 0)}")
            doc.add_paragraph(f"失败记录数: {summary.get('failed_records', 0)}")
            doc.add_paragraph(f"通过率: {summary.get('pass_rate', 0):.2%}")
            doc.add_paragraph("")  # 空行
            
            # 建议措施
            if data.recommendations:
                doc.add_heading("建议措施", level=1)
                for i, rec in enumerate(data.recommendations, 1):
                    doc.add_paragraph(f"{i}. {rec}")
                doc.add_paragraph("")  # 空行
                
            # 详细记录表格
            if self.config.include_details and not data.details.empty:
                doc.add_heading("详细记录", level=1)
                
                # 添加表格
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # 表头
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '记录ID'
                hdr_cells[1].text = '审核日期'
                hdr_cells[2].text = '任务类型'
                hdr_cells[3].text = '风险等级'
                hdr_cells[4].text = '状态'
                
                # 数据行
                for _, record in data.details.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(record.get('record_id', ''))
                    row_cells[1].text = str(record.get('audit_date', ''))
                    row_cells[2].text = str(record.get('task_type', ''))
                    row_cells[3].text = str(record.get('risk_level', ''))
                    row_cells[4].text = '通过' if record.get('passed', False) else '失败'
                    
            # 插入图表
            if self.config.include_charts and data.charts:
                doc.add_heading("图表分析", level=1)
                for chart_name, chart_path in data.charts.items():
                    if chart_path and Path(chart_path).exists():
                        doc.add_paragraph(f"{chart_name}:")
                        doc.add_picture(chart_path, width=Inches(6))
                        
            # 保存文档
            doc.save(report_path)
            
            logger.info(f"Word报告已生成: {report_path}")
            return str(report_path)
            
        except Exception as e:
            logger.error(f"生成Word报告失败: {e}")
            raise
            
    def generate_excel_report(self, data: ReportData) -> str:
        """生成Excel报告"""
        try:
            report_path = Path(self.config.output_dir) / f"audit_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            
            with pd.ExcelWriter(report_path, engine='openpyxl') as writer:
                # 概要表
                summary_df = pd.DataFrame([data.summary])
                summary_df.to_excel(writer, sheet_name='审核概要', index=False)
                
                # 详细记录表
                if not data.details.empty:
                    data.details.to_excel(writer, sheet_name='详细记录', index=False)
                    
                # 风险分析表
                if data.risk_analysis:
                    risk_df = pd.DataFrame([data.risk_analysis])
                    risk_df.to_excel(writer, sheet_name='风险分析', index=False)
                    
                # 建议措施表
                if data.recommendations:
                    rec_df = pd.DataFrame({
                        '序号': range(1, len(data.recommendations) + 1),
                        '建议措施': data.recommendations
                    })
                    rec_df.to_excel(writer, sheet_name='建议措施', index=False)
                    
            logger.info(f"Excel报告已生成: {report_path}")
            return str(report_path)
            
        except Exception as e:
            logger.error(f"生成Excel报告失败: {e}")
            raise
            
    def generate_json_report(self, data: ReportData) -> str:
        """生成JSON报告"""
        try:
            report_path = Path(self.config.output_dir) / f"audit_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            
            # 准备报告数据
            report_data = {
                'metadata': {
                    'title': self.config.title,
                    'author': self.config.author,
                    'generation_time': datetime.now().isoformat(),
                    'report_period': self.config.report_period,
                    'company': self.config.company
                },
                'summary': data.summary,
                'risk_analysis': data.risk_analysis,
                'trend_analysis': data.trend_analysis,
                'recommendations': data.recommendations,
                'charts': data.charts,
                'details': data.details.to_dict('records') if not data.details.empty else []
            }
            
            # 保存JSON文件
            with open(report_path, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2, default=str)
                
            logger.info(f"JSON报告已生成: {report_path}")
            return str(report_path)
            
        except Exception as e:
            logger.error(f"生成JSON报告失败: {e}")
            raise
            
    def generate_comprehensive_report(self, data: ReportData, formats: List[str] = None) -> Dict[str, str]:
        """生成综合报告（多种格式）"""
        if formats is None:
            formats = ['html', 'excel']
            
        # 生成图表
        if self.config.include_charts:
            self._generate_charts(data)
            
        generated_files = {}
        
        for format_type in formats:
            try:
                if format_type == 'html':
                    file_path = self.generate_html_report(data)
                elif format_type == 'pdf':
                    file_path = self.generate_pdf_report(data)
                elif format_type == 'word':
                    file_path = self.generate_word_report(data)
                elif format_type == 'excel':
                    file_path = self.generate_excel_report(data)
                elif format_type == 'json':
                    file_path = self.generate_json_report(data)
                else:
                    logger.warning(f"不支持的报告格式: {format_type}")
                    continue
                    
                generated_files[format_type] = file_path
                
            except Exception as e:
                logger.error(f"生成{format_type}格式报告失败: {e}")
                
        return generated_files
        
    def _generate_charts(self, data: ReportData):
        """生成所有图表"""
        try:
            # 风险分布图
            if data.risk_analysis:
                risk_chart = self.chart_generator.generate_risk_distribution_chart(data.risk_analysis)
                if risk_chart:
                    data.charts['risk_distribution'] = risk_chart
                    
            # 趋势分析图
            if data.trend_analysis:
                trend_chart = self.chart_generator.generate_trend_chart(data.trend_analysis)
                if trend_chart:
                    data.charts['trend_analysis'] = trend_chart
                    
            # 任务性能图
            if data.summary and 'task_distribution' in data.summary:
                task_chart = self.chart_generator.generate_task_performance_chart(data.summary)
                if task_chart:
                    data.charts['task_performance'] = task_chart
                    
        except Exception as e:
            logger.error(f"生成图表失败: {e}")


# 便捷函数
def generate_audit_report(audit_data: Dict[str, Any], 
                        formats: List[str] = ['html', 'excel'],
                        config: Optional[ReportConfig] = None) -> Dict[str, str]:
    """生成审核报告的便捷函数"""
    generator = ReportGenerator(config)
    
    # 准备报告数据
    report_data = ReportData(
        summary=audit_data.get('summary', {}),
        audit_results=audit_data.get('audit_results', []),
        risk_analysis=audit_data.get('risk_analysis', {}),
        trend_analysis=audit_data.get('trend_analysis', {}),
        recommendations=audit_data.get('recommendations', []),
        details=pd.DataFrame(audit_data.get('details', []))
    )
    
    return generator.generate_comprehensive_report(report_data, formats)
